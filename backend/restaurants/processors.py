import os
import json
import base64
import re
from datetime import datetime
from .dataset_builder import build_dataset

# ================================
# PATHS
# ================================
BASE_DIR        = os.path.dirname(__file__)
RESTAURANTS_DIR = BASE_DIR


# ================================================================
# CLAUDE VISION PARSER
# Handles ANY menu format — photo, PDF, Canva, Adobe, Excel, Word
# Sends the file to Claude as an image and extracts structured data
# ================================================================

def _ask_claude_vision(client, image_b64, media_type):
    """
    Sends one image to Claude Vision and returns structured menu items.
    Uses claude-opus for best accuracy on complex visual menus.
    """
    prompt = """You are reading a restaurant menu image. Extract ALL menu items visible.

For each item return a JSON object with:
- name: item name (required)
- category: menu section (e.g. Appetizers, Entrees, Desserts, Drinks, Sides)
- price: price as string like "12.99" with no dollar sign (empty string if not visible)
- description: item description (empty string if none)
- allergens: any allergen info (empty string if none)
- notes: any special notes (empty string if none)

Return ONLY a valid JSON array. No explanation, no markdown, no code blocks.
If you cannot read the menu, return [].

Example:
[{"name":"Caesar Salad","category":"Salads","price":"14.00","description":"Romaine, croutons, parmesan","allergens":"dairy, gluten","notes":""},
 {"name":"Grilled Salmon","category":"Entrees","price":"28.00","description":"With seasonal vegetables","allergens":"fish","notes":""}]"""

    try:
        response = client.messages.create(
            model="claude-opus-4-5",
            max_tokens=4096,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type":   "image",
                            "source": {
                                "type":       "base64",
                                "media_type": media_type,
                                "data":       image_b64,
                            }
                        },
                        {
                            "type": "text",
                            "text": prompt
                        }
                    ]
                }
            ]
        )

        raw = response.content[0].text.strip()
        raw = re.sub(r"```json|```", "", raw).strip()
        items = json.loads(raw)
        print(f"[vision] Extracted {len(items)} items")
        return items

    except json.JSONDecodeError as e:
        print(f"[vision] JSON parse error: {e}")
        return []
    except Exception as e:
        print(f"[vision] Claude vision error: {e}")
        return []


def parse_with_claude_vision(file_path):
    """
    Entry point for vision parsing.
    Converts file to base64 image and sends to Claude.
    Handles JPG, PNG, and PDF (converts pages to images).
    """
    import anthropic
    client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))

    ext = file_path.lower().split(".")[-1]

    # ---- JPG / PNG ----
    if ext in ("jpg", "jpeg"):
        with open(file_path, "rb") as f:
            data = base64.standard_b64encode(f.read()).decode("utf-8")
        return _ask_claude_vision(client, data, "image/jpeg")

    elif ext == "png":
        with open(file_path, "rb") as f:
            data = base64.standard_b64encode(f.read()).decode("utf-8")
        return _ask_claude_vision(client, data, "image/png")

    # ---- PDF — convert pages to images ----
    elif ext == "pdf":
        try:
            from pdf2image import convert_from_path
            import io
            pages   = convert_from_path(file_path, dpi=200, first_page=1, last_page=4)
            results = []
            for page_num, page in enumerate(pages):
                print(f"[vision] Processing PDF page {page_num + 1}")
                buf    = io.BytesIO()
                page.save(buf, format="JPEG")
                b64    = base64.standard_b64encode(buf.getvalue()).decode("utf-8")
                items  = _ask_claude_vision(client, b64, "image/jpeg")
                results.extend(items)
            return results
        except Exception as e:
            print(f"[vision] PDF error: {e} — falling back to text")
            return parse_with_text_fallback(file_path)

    return []


# ================================================================
# STRUCTURED FILE PARSER
# For Excel / CSV menus that are already organized in columns
# ================================================================

# ================================================================
# SMART COLUMN DETECTOR
# Scores each column header against known field types.
# Works regardless of exact column name — fuzzy keyword matching.
# ================================================================

# Keywords that signal each field type
COLUMN_SIGNALS = {
    "name": [
        "name", "item", "dish", "food", "product", "menu item",
        "item name", "food item", "title", "meal", "entree"
    ],
    "price": [
        "price", "cost", "rate", "amount", "charge", "fee",
        "price ($)", "menu price", "usd", "$"
    ],
    "description": [
        "description", "desc", "details", "about", "info",
        "ingredients", "notes about", "what", "content"
    ],
    "category": [
        "category", "section", "course", "type", "group",
        "menu section", "part", "department", "cat"
    ],
    "allergens": [
        "allergen", "allergy", "allergies", "contains", "dietary",
        "may contain", "allergy info", "food allergy", "allergen info",
        "intolerance", "free from", "gluten", "dairy", "nuts"
    ],
    "notes": [
        "note", "special", "modifier", "extra", "add on",
        "instruction", "comment", "remark", "flag"
    ],
}


def score_column(header, field):
    """Score how likely a column header matches a field type (0-100)."""
    h = str(header).lower().strip()
    for kw in COLUMN_SIGNALS[field]:
        if kw == h:
            return 100          # exact match
        if kw in h or h in kw:
            return 70           # partial match
    return 0


def detect_columns(headers):
    """
    Given a list of column headers, return a dict mapping
    field names to the best matching column header.
    Only maps a column if it scores above threshold.
    """
    mapping = {}
    used    = set()

    for field in ["name", "price", "description", "category", "allergens", "notes"]:
        best_score  = 30        # minimum threshold to count as a match
        best_header = None

        for h in headers:
            if h in used:
                continue
            score = score_column(h, field)
            if score > best_score:
                best_score  = score
                best_header = h

        if best_header:
            mapping[field] = best_header
            used.add(best_header)
            print(f"[smart_col] '{best_header}' → {field} (score {best_score})")

    return mapping


def get_cell(row, header, default=""):
    """Safely get a cell value, converting NaN to empty string."""
    if header is None:
        return default
    val = row.get(header, default)
    if val is None:
        return default
    s = str(val).strip()
    return "" if s.lower() in ("nan", "none", "n/a", "-") else s


def parse_structured_file(file_path):
    """
    Smart column detection — reads any Excel/CSV menu regardless
    of column header naming conventions.

    Step 1: Load all rows
    Step 2: Score each column header against field types
    Step 3: Map best matching columns to fields
    Step 4: Extract items using the detected mapping
    """
    items = []

    try:
        if file_path.endswith(".csv"):
            import csv
            with open(file_path, newline="", encoding="utf-8", errors="ignore") as f:
                rows = list(csv.DictReader(f))
        else:
            import pandas as pd
            df   = pd.read_excel(file_path)
            # Convert all column names to strings
            df.columns = [str(c) for c in df.columns]
            rows = df.to_dict(orient="records")

        if not rows:
            print("[structured] Empty file")
            return []

        # Get all column headers from first row
        headers = list(rows[0].keys())
        print(f"[structured] Columns found: {headers}")

        # Detect which column maps to which field
        col_map = detect_columns(headers)
        print(f"[structured] Column mapping: {col_map}")

        if "name" not in col_map:
            print("[structured] No name column detected — cannot parse as structured data")
            return []

        for row in rows:
            name = get_cell(row, col_map.get("name"))
            if not name:
                continue

            allergen_raw = get_cell(row, col_map.get("allergens"))

            items.append({
                "name":        name,
                "category":    get_cell(row, col_map.get("category")),
                "price":       get_cell(row, col_map.get("price")).replace("$", ""),
                "description": get_cell(row, col_map.get("description")),
                "allergens":   allergen_raw,
                "allergies":   allergen_raw,   # set both so merge always fires
                "notes":       get_cell(row, col_map.get("notes")),
            })

        print(f"[structured] Extracted {len(items)} items")

    except Exception as e:
        print(f"[structured] Error: {e}")

    return items


# ================================================================
# TEXT FALLBACK PARSER
# Used for Word docs and plain text files
# ================================================================

KNOWN_CATEGORY_KEYWORDS = {
    "appetizer","appetizers","entree","entrees",
    "soup","soups","salad","salads","dessert","desserts",
    "drink","drinks","beverage","beverages",
    "sandwich","sandwiches","breakfast","lunch","dinner",
    "side","sides","pizza","pasta","burger","burgers",
    "seafood","steak","grill","kids","kids menu",
    "special","specials","comfort","classic","classics","main","mains",
}

def is_category_line(line):
    stripped = line.strip()
    if len(stripped) > 40: return False
    if re.search(r"\$?\s*\d{1,3}(?:[.,]\d{2})?", stripped): return False
    lower = stripped.lower()
    for kw in KNOWN_CATEGORY_KEYWORDS:
        if lower == kw or lower.startswith(kw): return True
    if stripped == stripped.upper() and len(stripped) <= 25 and not any(c.isdigit() for c in stripped):
        return True
    return False


def extract_text_lines(file_path):
    lines = []
    if file_path.endswith(".docx"):
        from docx import Document
        doc = Document(file_path)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text: lines.append(" | ".join(row_text))
        for p in doc.paragraphs:
            if p.text.strip(): lines.append(p.text.strip())
    else:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            lines = f.readlines()
    return [l.strip() for l in lines if l.strip()]


def parse_menu_lines(lines):
    items            = []
    current_category = "Unknown"
    price_pattern    = re.compile(r"(?:\$?\s*\d{1,3}(?:[.,]\d{2})?)")
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if len(line) < 3: i += 1; continue
        if is_category_line(line):
            current_category = line.strip().title(); i += 1; continue
        match = price_pattern.search(line)
        if match or (any(c.isalpha() for c in line) and i + 1 < len(lines)):
            price = ""; name = line
            if match:
                price = match.group().replace("$", "").strip()
                name  = line.replace(match.group(), "").strip()
            description = ""; j = i + 1
            while j < len(lines):
                next_line = lines[j].strip()
                if len(next_line) < 3: j += 1; continue
                if price_pattern.search(next_line): break
                if is_category_line(next_line): break
                description += next_line + " "; j += 1
            name = name.strip()
            if name and len(name) >= 3:
                items.append({"name": name, "category": current_category,
                              "price": price, "description": description.strip(),
                              "allergens": "", "notes": ""})
            i = j; continue
        i += 1
    return items


def parse_with_text_fallback(file_path):
    lines = extract_text_lines(file_path)
    return parse_menu_lines(lines) if lines else []


# ================================================================
# SMART DISPATCHER
# Picks the right parser based on file type — no guessing
# ================================================================

def extract_pdf_text_lines(file_path):
    """Extract text from PDF text layer using pdfplumber."""
    try:
        import pdfplumber
        lines = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    for line in text.split(chr(10)):
                        line = line.strip()
                        if line:
                            lines.append(line)
        print(f"[pdf] pdfplumber extracted {len(lines)} lines")
        return lines
    except ImportError:
        print("[pdf] pdfplumber not installed")
        return []
    except Exception as e:
        print(f"[pdf] pdfplumber error: {e}")
        return []


def pdf_to_images_and_parse(file_path):
    """Convert PDF pages to images then send to Claude Vision."""
    try:
        from pdf2image import convert_from_path
        import io
        import anthropic
        client  = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
        pages   = convert_from_path(file_path, dpi=200, first_page=1, last_page=4)
        results = []
        for page_num, page in enumerate(pages):
            print(f"[pdf] Page {page_num + 1} → Claude Vision")
            buf = io.BytesIO()
            page.save(buf, format="JPEG")
            b64 = base64.standard_b64encode(buf.getvalue()).decode("utf-8")
            items = _ask_claude_vision(client, b64, "image/jpeg")
            results.extend(items)
        return results
    except ImportError:
        print("[pdf] pdf2image/poppler not installed")
        return []
    except Exception as e:
        print(f"[pdf] pdf2image error: {e}")
        return []


def smart_parse(file_path):
    """
    ULI + Claude Vision two-layer parser.

    Layer 1 — ULI spatial parser (fast, free, no API call)
      Runs on text-extractable files: DOCX, TXT, CSV, XLSX
      Uses price anchors + spatial relationships + typography signals
      Reports a confidence score after parsing

    Layer 2 — Claude Vision (accurate, handles any visual layout)
      Runs on images and PDFs always (no text to extract)
      Also runs as fallback if ULI confidence score is too low

    This means:
    - Simple text menus   → ULI handles it, no API cost
    - Complex photo menus → Claude Vision handles it
    - Word/Excel menus    → ULI tries first, Vision if low confidence
    """
    from .uli_parser import uli_parse_file, is_garbage_text

    ext = file_path.lower().split(".")[-1]
    print(f"\n[smart_parse] {os.path.basename(file_path)} — type: {ext}")

    # IMAGES — Claude Vision always
    if ext in ("jpg", "jpeg", "png"):
        print("[smart_parse] Image → Claude Vision")
        return parse_with_claude_vision(file_path)

    # PDF — 3 layer approach
    elif ext == "pdf":
        print("[smart_parse] PDF — trying 3 layers")

        # Layer 1: pdfplumber text extraction
        lines = extract_pdf_text_lines(file_path)
        if lines and not is_garbage_text(lines):
            items, report = uli_parse_file(lines)
            if not report["fallback"]:
                print(f"[smart_parse] PDF text → ULI success (score {report['score']})")
                return items
            print(f"[smart_parse] PDF text low confidence → trying image conversion")
        else:
            print("[smart_parse] No usable text in PDF → trying image conversion")

        # Layer 2: pdf2image + poppler
        results = pdf_to_images_and_parse(file_path)
        if results:
            print(f"[smart_parse] PDF images → got {len(results)} items")
            return results

        # Layer 3: Claude Vision direct fallback
        print("[smart_parse] PDF last resort → Claude Vision direct")
        return parse_with_claude_vision(file_path)

    # EXCEL / CSV — structured parser
    elif ext in ("xlsx", "xls", "csv"):
        print("[smart_parse] Spreadsheet → Structured parser")
        items = parse_structured_file(file_path)
        if items:
            return items
        print("[smart_parse] No columns → Claude Vision fallback")
        return parse_with_claude_vision(file_path)

    # WORD / TXT — ULI first, Vision fallback
    elif ext in ("docx", "doc", "txt"):
        print("[smart_parse] Text file → ULI")
        lines = extract_text_lines(file_path)
        if not lines:
            print("[smart_parse] No text → Claude Vision fallback")
            return parse_with_claude_vision(file_path)
        items, report = uli_parse_file(lines)
        if report["fallback"]:
            print(f"[smart_parse] ULI score {report['score']} low → Claude Vision fallback")
            return parse_with_claude_vision(file_path)
        print(f"[smart_parse] ULI confident (score {report['score']})")
        return items

    # UNKNOWN — ULI text fallback
    else:
        print("[smart_parse] Unknown type → ULI text fallback")
        lines = extract_text_lines(file_path)
        if not lines:
            return []
        items, report = uli_parse_file(lines)
        return items


# ================================================================
# DATASET STRUCTURE
# ================================================================

def ensure_dataset_structure(restaurant_id):
    dataset_path = os.path.join(BASE_DIR, restaurant_id, "menu_dataset.json")
    os.makedirs(os.path.dirname(dataset_path), exist_ok=True)
    if not os.path.exists(dataset_path):
        with open(dataset_path, "w", encoding="utf-8") as f:
            json.dump({"items": [], "categories": []}, f, indent=2)
        return
    with open(dataset_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    changed = False
    if isinstance(data, list):
        data = {"items": data, "categories": []}; changed = True
    if "categories" not in data:
        data["categories"] = []; changed = True
    if changed:
        with open(dataset_path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2)


# ================================================================
# FILE REGISTRATION
# ================================================================

def process_uploaded_file(restaurant_id, filename, category):
    files_index = os.path.join(BASE_DIR, restaurant_id, "files", "files.json")
    if not os.path.exists(files_index):
        os.makedirs(os.path.dirname(files_index), exist_ok=True)
        with open(files_index, "w", encoding="utf-8") as f:
            json.dump({"files": []}, f, indent=2)
    with open(files_index, "r", encoding="utf-8") as f:
        data = json.load(f)
    record = {
        "restaurant_id": restaurant_id,
        "name":          filename,
        "category":      category,
        "processed":     False,
        "uploaded_at":   datetime.utcnow().isoformat()
    }
    data["files"].append(record)
    with open(files_index, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2)
    return record


# ================================================================
# MAIN PROCESSOR | MRK | ULI | PIE AI
# ================================================================

def dispatch_unprocessed_files():

    from .dataset_builder import heal_existing_dataset

    for restaurant_id in os.listdir(BASE_DIR):

        restaurant_path = os.path.join(BASE_DIR, restaurant_id)
        if not os.path.isdir(restaurant_path):
            continue

        files_index = os.path.join(BASE_DIR, restaurant_id, "files", "files.json")

        heal_existing_dataset(restaurant_id)

        if not os.path.exists(files_index):
            continue

        with open(files_index, "r", encoding="utf-8") as f:
            data = json.load(f)

        print(f"\nPROCESSING: {restaurant_id}")

        for record in data.get("files", []):

            if record.get("processed"):
                continue

            ensure_dataset_structure(restaurant_id)

            file_path = os.path.join(
                BASE_DIR, restaurant_id, "files",
                record["category"], record["name"]
            )

            if not os.path.exists(file_path):
                print(f"Missing: {file_path}")
                continue

            print(f"\n→ {record['name']}")

            parsed_items = smart_parse(file_path)

            print(f"  Got {len(parsed_items)} items")

            if parsed_items:
                for item in parsed_items[:5]:
                    print(f"  • {item.get('name')} | {item.get('category')} | ${item.get('price')}")

            if not parsed_items:
                print("  No items — marking processed")
                mark_as_processed(record)
                continue

            build_dataset(record, parsed_items)
            mark_as_processed(record)


# ================================================================
# MARK PROCESSED
# ================================================================

def mark_as_processed(record):
    restaurant_id = record.get("restaurant_id")
    files_index   = os.path.join(BASE_DIR, restaurant_id, "files", "files.json")
    if not os.path.exists(files_index):
        return
    with open(files_index, "r", encoding="utf-8") as f:
        data = json.load(f)
    for f_record in data.get("files", []):
        if f_record["name"] == record["name"] and f_record["restaurant_id"] == record["restaurant_id"]:
            f_record["processed"] = True
    with open(files_index, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2)